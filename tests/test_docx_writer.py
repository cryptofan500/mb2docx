"""Tests for docx_writer.py — structure-level round-trip assertions.

Build a Document via the writer, optionally save and re-open it with
python-docx, and assert the resulting paragraphs / runs / paragraph
formats match the forensic expectations encoded in ``DocxStyleConfig``.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from mb2docx.config import CL_STYLE, CV_STYLE
from mb2docx.docx_writer import (
    _parse_bold_markup,
    build_docx,
    safe_save_docx,
)
from mb2docx.model import (
    ContactBlock,
    HeadingBlock,
    JobEntryBlock,
    ListBlock,
    SectionHeadingBlock,
)


def _first_paragraph_with_text(doc, predicate):
    """Return the first paragraph whose text matches the predicate."""
    for p in doc.paragraphs:
        if predicate(p.text):
            return p
    raise AssertionError("No paragraph matched the predicate")


def test_cover_letter_name_18pt_bold_centered():
    """Cover letter name: 18pt, bold, centered."""
    blocks = [
        HeadingBlock(type="heading", level=1, text="JANE DOE"),
        ContactBlock(type="contact_header", text="email@test.com | (555) 123-4567"),
    ]
    doc = build_docx(blocks, cfg=CL_STYLE)
    name_p = _first_paragraph_with_text(doc, lambda t: "JANE DOE" in t)

    assert name_p.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert len(name_p.runs) == 1
    run = name_p.runs[0]
    assert run.bold is True
    assert run.font.size == Pt(18)


def test_cv_section_heading_12pt_bold_all_caps():
    """CV section heading: 12pt, bold, ALL CAPS."""
    blocks = [SectionHeadingBlock(type="section_heading", text="Education")]
    doc = build_docx(blocks, cfg=CV_STYLE)
    p = doc.paragraphs[0]

    # Writer uppercases the rendered text regardless of input case.
    assert p.text == "EDUCATION"
    assert len(p.runs) == 1
    run = p.runs[0]
    assert run.bold is True
    assert run.font.size == Pt(12)


def test_job_entry_pipe_separator_title_bold_date_not_bold():
    """Job entry with PIPE separator: title bold, date not bold, ' | ' between."""
    blocks = [
        JobEntryBlock(
            type="job_entry",
            title="Senior Manager",
            date_range="January 2020 - Present",
        )
    ]
    doc = build_docx(blocks, cfg=CV_STYLE)  # CV_STYLE has separator="PIPE"
    p = doc.paragraphs[0]

    # Two runs: one bold title, one non-bold separator+date.
    assert len(p.runs) == 2
    title_run, date_run = p.runs
    assert title_run.text == "Senior Manager"
    assert title_run.bold is True
    assert date_run.text == " | January 2020 - Present"
    # Date must NOT be bold. Writer never explicitly sets bold=False, so it
    # comes back as None; the contract is "not truthy".
    assert not date_run.bold
    # Sanity: full text contains the pipe separator.
    assert "Senior Manager | January 2020 - Present" == p.text


def test_cv_bullet_has_hanging_indent():
    """CV bullet paragraphs use left=0.5", first_line_indent=-0.25"."""
    blocks = [
        ListBlock(type="list", ordered=False, items=["First item", "Second item"])
    ]
    doc = build_docx(blocks, cfg=CV_STYLE)

    bullet_paragraphs = [p for p in doc.paragraphs if p.text.startswith("•")]
    assert len(bullet_paragraphs) == 2
    for p in bullet_paragraphs:
        pf = p.paragraph_format
        assert pf.left_indent == Inches(0.5)
        assert pf.first_line_indent == Inches(-0.25)


def test_safe_save_docx_round_trip(tmp_path: Path):
    """safe_save_docx writes a file that python-docx can re-open with content."""
    blocks = [
        HeadingBlock(type="heading", level=1, text="JANE DOE"),
        ContactBlock(type="contact_header", text="email@test.com | (555) 123-4567"),
        SectionHeadingBlock(type="section_heading", text="EDUCATION"),
    ]
    doc = build_docx(blocks, cfg=CV_STYLE)
    out_path = tmp_path / "round_trip.docx"

    result = safe_save_docx(doc, out_path)
    assert result.path == out_path
    assert out_path.exists()

    reopened = Document(str(out_path))
    assert len(reopened.paragraphs) > 0
    texts = [p.text for p in reopened.paragraphs]
    assert any("JANE DOE" in t for t in texts)
    assert any("EDUCATION" in t for t in texts)


def test_parse_bold_markup_splits_runs_correctly():
    """_parse_bold_markup turns **bold** into the right (text, is_bold) tuples."""
    assert _parse_bold_markup("plain text") == [("plain text", False)]

    parts = _parse_bold_markup("regular **bold** more")
    assert parts == [
        ("regular ", False),
        ("bold", True),
        (" more", False),
    ]

    # Multiple bolds.
    parts = _parse_bold_markup("**A** middle **B**")
    assert parts == [
        ("A", True),
        (" middle ", False),
        ("B", True),
    ]

    # Bold at the very start with no trailing text.
    parts = _parse_bold_markup("**only**")
    assert parts == [("only", True)]


def test_cv_example_round_trip_uses_real_input(tmp_path: Path):
    """Bundled examples/cv.md parses and renders to a saveable DOCX."""
    from mb2docx.parser import parse_markdown_like

    cv_md = Path(__file__).resolve().parent.parent / "examples" / "cv.md"
    text = cv_md.read_text(encoding="utf-8")
    blocks = parse_markdown_like(text, is_cover_letter=False)
    doc = build_docx(blocks, cfg=CV_STYLE)

    out_path = tmp_path / "cv_round_trip.docx"
    safe_save_docx(doc, out_path)

    reopened = Document(str(out_path))
    paragraphs = [p.text for p in reopened.paragraphs]
    assert any("JANE DOE" in t for t in paragraphs)
    assert any("PROFESSIONAL SUMMARY" in t for t in paragraphs)
